from flask import Flask, render_template, request, jsonify
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain.chains.summarize import load_summarize_chain
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.docstore.document import Document
from dotenv import load_dotenv
import asyncio
import docx
import pandas as pd
from urllib.request import urlopen
from bs4 import BeautifulSoup
import certifi
import ssl

# Load environment variables
load_dotenv()

# Retrieve environment variables
google_api_key = os.getenv("GOOGLE_API_KEY")

# Check if environment variables are loaded correctly
if not google_api_key:
    raise ValueError("GOOGLE_API_KEY is not set correctly")

# Configure Google Generative AI
genai.configure(api_key=google_api_key)

# Define the directory to save uploaded files
UPLOAD_DIR = "uploaded_files"
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    chunks = text_splitter.split_text(text)
    docs = [Document(page_content=chunk) for chunk in chunks]
    return docs

def get_summary_chain():
    model = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.7)
    chain = load_summarize_chain(model,chain_type="map_reduce")
    return chain


async def summarize_text(text):
    chain = get_summary_chain()
    docs = get_text_chunks(text)
    summary = chain.run(docs)

    return summary.strip()


def extract_text_from_pdf(pdf_path):
    text = ""
    pdf_reader = PdfReader(pdf_path)
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_text_from_csv(csv_path):
    df = pd.read_csv(csv_path)
    return df.to_string(index=False)

def extract_text_from_txt(txt_path):
    with open(txt_path, 'r', encoding='utf-8') as file:
        return file.read()

def extract_text_from_excel(excel_path):
    df = pd.read_excel(excel_path)
    return df.to_string(index=False)

def extract_text_from_url(url):
    context = ssl.create_default_context(cafile=certifi.where())
    with urlopen(url, context=context) as response:
        html = response.read()
    soup = BeautifulSoup(html, 'html.parser')
    return soup.get_text()

app = Flask(__name__)

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload():
    if len(request.files) == 0:
        return jsonify({"error": "No file uploaded"}), 400

    uploaded_file = None
    for key in request.files:
        uploaded_file = request.files[key]
        break  # Get the first file found

    if uploaded_file is None or uploaded_file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    file_path = os.path.join(UPLOAD_DIR, uploaded_file.filename)
    uploaded_file.save(file_path)

    file_extension = uploaded_file.filename.split('.')[-1].lower()
    text = ""

    if file_extension == "pdf":
        text = extract_text_from_pdf(file_path)
    elif file_extension == "docx":
        text = extract_text_from_docx(file_path)
    elif file_extension == "csv":
        text = extract_text_from_csv(file_path)
    elif file_extension == "txt":
        text = extract_text_from_txt(file_path)
    elif file_extension in ["xls", "xlsx"]:
        text = extract_text_from_excel(file_path)
    else:
        return jsonify({"error": "Unsupported file type"}), 400

    summary = asyncio.run(summarize_text(text))

    return jsonify({"summary": summary}), 200

@app.route("/upload_url", methods=["POST"])
def upload_url():
    data = request.json
    url = data.get('url')
    if not url:
        return jsonify({"error": "No URL provided"}), 400

    try:
        text = extract_text_from_url(url)
    except Exception as e:
        return jsonify({"error": f"Failed to extract text from URL: {str(e)}"}), 400

    summary = asyncio.run(summarize_text(text))

    return jsonify({"summary": summary}), 200

if __name__ == "__main__":
    app.run(debug=True)
